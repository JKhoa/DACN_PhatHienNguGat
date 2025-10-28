import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.dirname(os.path.dirname(__file__))
DOCS_DIR = os.path.join(ROOT, 'docs')
OUT_PATH = os.path.join(DOCS_DIR, 'TIEU_LUAN_DECORATOR.docx')


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(12)
    return p


def add_code_block(doc: Document, code: str, language_hint: str = ''):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    if language_hint:
        p = doc.add_paragraph()
        p.add_run(f"(Ngôn ngữ: {language_hint})").italic = True


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TIỂU LUẬN: DECORATOR PATTERN\nTRONG ỨNG DỤNG PHÁT HIỆN NGỦ GẬT TRONG LỚP HỌC')
    r.bold = True
    r.font.size = Pt(18)

    p2 = doc.add_paragraph('\n')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Dự án: DACN_PhatHienNguGat\nNgày: 21/10/2025')
    r2.font.size = Pt(12)


def add_toc_placeholder(doc: Document):
    add_heading(doc, 'Mục lục', level=1)
    add_paragraph(doc, 'Ghi chú: Hãy vào References > Table of Contents trong Word và Insert TOC để tạo mục lục tự động rồi nhấn F9 để cập nhật.', italic=True)


def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    doc = Document()

    # Title page
    add_title_page(doc)

    # TOC placeholder
    add_toc_placeholder(doc)

    # I. Giới thiệu
    add_heading(doc, 'I. Giới thiệu', level=1)
    add_paragraph(doc, 'Bối cảnh lớp học hiện đại yêu cầu các hệ thống hỗ trợ giảng dạy có khả năng theo dõi trạng thái chú ý của học viên. '
                      'Ứng dụng phát hiện ngủ gật sử dụng nhiều camera đòi hỏi độ trễ thấp và tốc độ khung hình ổn định để kịp thời cảnh báo và can thiệp.')
    add_paragraph(doc, 'Mẫu thiết kế Decorator được lựa chọn nhằm tách rời các mối quan tâm (capture, đo hiệu năng, suy luận, overlay, ghi log, reconnect) '
                      'và cho phép mở rộng linh hoạt mà không thay đổi thành phần lõi.')

    # II. Cơ sở lý thuyết Decorator
    add_heading(doc, 'II. Cơ sở lý thuyết Decorator', level=1)
    add_paragraph(doc, 'Decorator thuộc nhóm cấu trúc (Structural) trong GoF. Ý tưởng chính là bọc (wrap) một đối tượng để bổ sung hành vi động, '
                      'không sửa đổi lớp gốc. Điều này phù hợp với pipeline video nơi mỗi bước xử lý có thể bật/tắt theo cấu hình.')

    # UML snippets as code (PlantUML)
    add_heading(doc, 'Sơ đồ lớp (PlantUML)', level=2)
    class_uml = """@startuml
interface ICameraStream {
  +read(): Frame
}

class WebcamStream implements ICameraStream
class IPCameraStream implements ICameraStream

abstract class StreamDecorator implements ICameraStream {
  -inner: ICameraStream
  +read(): Frame
}

class FrameQueueDecorator extends StreamDecorator
class DetectionDecorator extends StreamDecorator
class OverlayDecorator extends StreamDecorator
class LoggingDecorator extends StreamDecorator
class PerformanceDecorator extends StreamDecorator
class ReconnectDecorator extends StreamDecorator

ICameraStream <|.. WebcamStream
ICameraStream <|.. IPCameraStream
ICameraStream <|.. StreamDecorator
StreamDecorator <|-- FrameQueueDecorator
StreamDecorator <|-- DetectionDecorator
StreamDecorator <|-- OverlayDecorator
StreamDecorator <|-- LoggingDecorator
StreamDecorator <|-- PerformanceDecorator
StreamDecorator <|-- ReconnectDecorator

StreamDecorator o-- ICameraStream : wraps
@enduml"""
    add_code_block(doc, class_uml, 'plantuml')

    add_heading(doc, 'Sơ đồ trình tự (PlantUML)', level=2)
    seq_uml = """@startuml
actor User
participant UI
participant "Stream (stacked decorators)" as S
participant Detector
participant Logger

User -> UI: Start streaming
UI -> S: read()
S -> S: FrameQueueDecorator.read()
S -> S: DetectionDecorator.read()
S -> Detector: infer(frame)
Detector --> S: result
S -> S: OverlayDecorator.read()
S -> Logger: log(event)
S --> UI: frame+overlay
UI -> User: display
@enduml"""
    add_code_block(doc, seq_uml, 'plantuml')

    # III. Phân tích hệ thống áp dụng
    add_heading(doc, 'III. Phân tích hệ thống áp dụng', level=1)
    add_paragraph(doc, 'Dự án sử dụng Ultralytics YOLO pose (yolo11n/11s) kết hợp state machine buồn ngủ/thức dậy. '
                      'Luồng xử lý được tách rời thành các decorator độc lập, giảm coupling và tăng khả năng kiểm thử.')

    # IV. Thiết kế & triển khai
    add_heading(doc, 'IV. Thiết kế & triển khai', level=1)
    add_heading(doc, 'Trích đoạn mã minh họa Decorator (Python)', level=2)
    decorator_code = """from abc import ABC, abstractmethod
from typing import Optional, Any
import time

class ICameraStream(ABC):
    @abstractmethod
    def read(self) -> Optional[Any]:
        ...

class WebcamStream(ICameraStream):
    def __init__(self, device_id: int = 0):
        import cv2
        self.cap = cv2.VideoCapture(device_id)
        self.cap.set(cv2.CAP_PROP_BUFFERSIZE, 1)
    def read(self):
        ok, frame = self.cap.read()
        return frame if ok else None

class StreamDecorator(ICameraStream):
    def __init__(self, inner: ICameraStream):
        self.inner = inner
    def read(self):
        return self.inner.read()

class FrameQueueDecorator(StreamDecorator):
    def __init__(self, inner: ICameraStream, maxsize: int = 3):
        super().__init__(inner)
        from queue import Queue
        import threading
        self.q = Queue(maxsize=maxsize)
        self.running = True
        def loop():
            while self.running:
                f = self.inner.read()
                if f is None:
                    time.sleep(0.02); continue
                if self.q.full():
                    try: self.q.get_nowait()
                    except: pass
                self.q.put_nowait(f)
        self.th = threading.Thread(target=loop, daemon=True)
        self.th.start()
    def read(self):
        f = None
        while not self.q.empty():
            f = self.q.get()
        return f

class PerformanceDecorator(StreamDecorator):
    def __init__(self, inner: ICameraStream):
        super().__init__(inner)
        self.last = time.time(); self.fps = 0.0
    def read(self):
        f = self.inner.read()
        now = time.time(); dt = max(now - self.last, 1e-6)
        self.fps = 1.0/dt
        self.last = now
        return f"""
    add_code_block(doc, decorator_code, 'python')

    add_heading(doc, 'Cấu hình YAML cho lớp học', level=2)
    yaml_example = """cameras:
  - name: FrontRow
    type: webcam
    device_id: 0
    decorators:
      - FrameQueueDecorator: { maxsize: 3 }
      - PerformanceDecorator: {}
      - DetectionDecorator: { model: yolo11n-pose.pt, conf: 0.25, strategy: yolo }
      - OverlayDecorator: { show_fps: true, show_state: true }
      - LoggingDecorator: { path: logs/frontrow_events.csv }
  - name: BackRow
    type: ipcam
    brand: hikvision
    ip: 192.168.1.21
    port: 554
    username: admin
    password: '***'
    stream_quality: sub
    decorators:
      - ReconnectDecorator: { retries: 5, backoff_s: 2 }
      - FrameQueueDecorator: { maxsize: 5 }
      - PerformanceDecorator: {}
      - DetectionDecorator: { model: yolo11s-pose.pt, conf: 0.3, strategy: yolo }
      - OverlayDecorator: { show_fps: true, show_state: true }
      - LoggingDecorator: { path: logs/backrow_events.csv }
  - name: TeacherSide
    type: ipcam
    brand: dahua
    ip: 192.168.1.22
    port: 554
    username: admin
    password: '***'
    stream_quality: main
    decorators:
      - FrameQueueDecorator: { maxsize: 3 }
      - PerformanceDecorator: {}
      - DetectionDecorator: { model: yolo11n-pose.pt, conf: 0.25, strategy: mediapipe }
      - OverlayDecorator: { show_fps: false, show_state: true }"""
    add_code_block(doc, yaml_example, 'yaml')

    # V. Đánh giá & so sánh
    add_heading(doc, 'V. Đánh giá & so sánh', level=1)
    add_paragraph(doc, 'Đề xuất benchmark trong bối cảnh phòng học: 1/4/8 camera, 480p/720p/1080p, so sánh CPU-only vs GPU. '
                      'Tiêu chí đạt: FPS ≥ 30 và độ trễ ≤ 2s. Kết quả được tổng hợp dưới dạng bảng và biểu đồ.')

    # VI. Kết luận & hướng phát triển
    add_heading(doc, 'VI. Kết luận & hướng phát triển', level=1)
    add_paragraph(doc, 'Decorator đem lại mô-đun hóa cao, khả năng mở rộng và kiểm thử tốt cho ứng dụng phát hiện ngủ gật trong lớp học. '
                      'Hướng phát triển gồm tích hợp Strategy cho các thuật toán phát hiện, bổ sung PrivacyMaskDecorator và tối ưu pipeline đa GPU.')

    # Appendix: Full prompt content
    prompt_path = os.path.join(DOCS_DIR, 'TIEU_LUAN_DECORATOR_PROMPT.txt')
    if os.path.exists(prompt_path):
        add_heading(doc, 'Phụ lục: Prompt đầy đủ để sinh báo cáo', level=1)
        try:
            with open(prompt_path, 'r', encoding='utf-8') as f:
                prompt_text = f.read()
            # Split into chunks to avoid extremely long runs
            for chunk in prompt_text.split('\n'):
                add_code_block(doc, chunk, '')
        except Exception as e:
            add_paragraph(doc, f'Không thể đọc prompt: {e}')

    # Save
    doc.save(OUT_PATH)
    print(f"Report generated: {OUT_PATH}")


if __name__ == '__main__':
    main()
